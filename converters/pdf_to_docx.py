"""PDF -> DOCX converter that produces a clean, editable, reflowable document.

Unlike pixel-perfect engines (e.g. pdf2docx) that wrap every block in absolutely
positioned frames -- which causes blank pages and broken layouts, especially with
two-column academic papers -- this rebuilds the text as normal flowing Word
paragraphs in proper reading order. Optimised for papers and ordinary documents.
"""
from io import BytesIO
import re
import statistics

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches

# PyMuPDF span flag bits
_FLAG_ITALIC = 1 << 1
_FLAG_BOLD = 1 << 4

_CONTENT_WIDTH_IN = 6.0  # usable width for images on a Letter/A4 page with margins


def convert(input_path: str, output_path: str) -> None:
    """Convert PDF to a clean, editable DOCX. Raises on failure."""
    pdf = fitz.open(input_path)
    try:
        doc = Document()
        body_size = _estimate_body_size(pdf)
        if body_size:
            doc.styles["Normal"].font.size = Pt(max(9, min(12, round(body_size))))

        running = _running_headers(pdf)
        for page in pdf:
            blocks = _ordered_blocks(page, running)
            for block in blocks:
                if block["type"] == 1:
                    _add_image(doc, block)
                else:
                    _add_text_block(doc, block, body_size)
        doc.save(output_path)
    finally:
        pdf.close()


def _estimate_body_size(pdf) -> float:
    """Most common span font size across the document = body text size."""
    sizes = []
    for page in pdf:
        for block in page.get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("text", "").strip():
                        sizes.append(round(span["size"], 1))
    if not sizes:
        return 0.0
    try:
        return statistics.mode(sizes)
    except statistics.StatisticsError:
        return statistics.median(sizes)


def _norm_sig(text: str) -> str:
    """Normalise text for matching running headers/footers (page numbers vary)."""
    return re.sub(r"\s+", " ", re.sub(r"\d+", "", text)).strip().lower()


def _running_headers(pdf) -> set:
    """Signatures of text that repeats across pages in the top/bottom margins."""
    counts = {}
    npages = pdf.page_count
    for page in pdf:
        ph = page.rect.height
        seen = set()
        for b in page.get_text("dict").get("blocks", []):
            if b.get("type") != 0:
                continue
            y0, y1 = b["bbox"][1], b["bbox"][3]
            if not (y1 < 0.12 * ph or y0 > 0.88 * ph):
                continue
            text = _block_text(b).strip()
            if not text or len(text) > 90:
                continue
            sig = _norm_sig(text)
            if len(sig) >= 4:  # skip pure numbers / too-generic fragments
                seen.add(sig)
        for sig in seen:
            counts[sig] = counts.get(sig, 0) + 1
    threshold = max(3, npages * 0.3)
    return {sig for sig, c in counts.items() if c >= threshold}


def _ordered_blocks(page, running=frozenset()):
    """Return page blocks in reading order, dropping headers/footers.

    Detects a two-column layout and reads the left column top-to-bottom, then the
    right column. Full-width blocks (titles, abstracts) are kept as anchors.
    """
    rect = page.rect
    pw, ph = rect.width, rect.height
    raw = page.get_text("dict").get("blocks", [])

    blocks = []
    for b in raw:
        x0, y0, x1, y1 = b["bbox"]
        if b.get("type") == 1:  # image
            blocks.append(b)
            continue
        text = _block_text(b).strip()
        if not text:
            continue
        in_margin = y1 < 0.12 * ph or y0 > 0.88 * ph
        # Drop short, single-line page numbers right at the top/bottom edge.
        if ((y1 < 0.04 * ph or y0 > 0.96 * ph)
                and len(text) < 70 and len(b.get("lines", [])) <= 1):
            continue
        # Drop running headers/footers that repeat across pages.
        if in_margin and _norm_sig(text) in running:
            continue
        blocks.append(b)

    if not blocks:
        return []

    mid = pw / 2
    # Full-width blocks (titles, abstracts, spanning figures) are normal in a
    # two-column layout, so only consider *narrow* blocks when deciding columns.
    narrow = [b for b in blocks
              if b.get("type") != 1 and (b["bbox"][2] - b["bbox"][0]) <= 0.6 * pw]
    left = sum(1 for b in narrow if b["bbox"][2] < mid)
    right = sum(1 for b in narrow if b["bbox"][0] > mid)
    straddle = len(narrow) - left - right
    two_col = left >= 2 and right >= 2 and straddle <= 0.25 * (left + right)

    if not two_col:
        return sorted(blocks, key=lambda b: (round(b["bbox"][1]), b["bbox"][0]))

    # Band-based ordering: full-width blocks (titles, abstracts, spanning
    # figures) act as horizontal separators. Within each band between them, read
    # the left column fully, then the right column.
    def is_anchor(b):
        x0, _, x1, _ = b["bbox"]
        w = x1 - x0
        straddles = x0 < mid < x1
        return w > 0.6 * pw or (straddles and w > 0.4 * pw)

    anchors, col_l, col_r = [], [], []
    for b in blocks:
        if is_anchor(b):
            anchors.append(b)
        elif (b["bbox"][0] + b["bbox"][2]) / 2 < mid:
            col_l.append(b)
        else:
            col_r.append(b)

    anchors.sort(key=lambda b: b["bbox"][1])
    col_l.sort(key=lambda b: b["bbox"][1])
    col_r.sort(key=lambda b: b["bbox"][1])

    result = []
    band_top = float("-inf")
    for anchor in anchors + [None]:
        band_bottom = anchor["bbox"][1] if anchor else float("inf")
        band_l = [b for b in col_l if band_top <= b["bbox"][1] < band_bottom]
        band_r = [b for b in col_r if band_top <= b["bbox"][1] < band_bottom]
        result.extend(band_l)
        result.extend(band_r)
        if anchor is not None:
            result.append(anchor)
            band_top = anchor["bbox"][1]
    return result


def _block_text(block) -> str:
    return "".join(
        span.get("text", "")
        for line in block.get("lines", [])
        for span in line.get("spans", [])
    )


def _add_text_block(doc, block, body_size):
    """Emit a block as a heading or a body paragraph, joining hyphenated lines."""
    runs = []  # list of [text, bold, italic]
    max_size = 0.0
    for line in block.get("lines", []):
        line_runs = []
        for span in line.get("spans", []):
            text = span.get("text", "")
            if not text:
                continue
            max_size = max(max_size, span.get("size", 0))
            line_runs.append([
                text,
                bool(span.get("flags", 0) & _FLAG_BOLD),
                bool(span.get("flags", 0) & _FLAG_ITALIC),
            ])
        if not line_runs:
            continue
        if runs:  # join with previous line
            prev = runs[-1][0].rstrip()
            if prev.endswith("-"):
                runs[-1][0] = prev[:-1]  # de-hyphenate
            else:
                runs[-1][0] = prev + " "
        runs.extend(line_runs)

    if not runs:
        return

    full_text = "".join(r[0] for r in runs).strip()
    if not full_text:
        return

    # Heading detection: short, single-paragraph, noticeably larger than body.
    level = 0
    if body_size and len(block.get("lines", [])) <= 2 and len(full_text) < 120:
        ratio = max_size / body_size if body_size else 1
        if ratio >= 1.8:
            level = 1
        elif ratio >= 1.45:
            level = 2
        elif ratio >= 1.22:
            level = 3

    if level:
        doc.add_heading(full_text, level=level)
        return

    para = doc.add_paragraph()
    for text, bold, italic in runs:
        if not text:
            continue
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic


def _add_image(doc, block):
    try:
        img = block.get("image")
        if not img:
            return
        x0, _, x1, _ = block["bbox"]
        width_in = min(_CONTENT_WIDTH_IN, max(1.0, (x1 - x0) / 72.0))
        doc.add_picture(BytesIO(img), width=Inches(width_in))
    except Exception:
        pass  # skip un-embeddable images rather than fail the whole conversion
