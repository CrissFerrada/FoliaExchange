import pytest
from fpdf import FPDF
from fpdf.enums import XPos, YPos
from docx import Document


@pytest.fixture
def sample_pdf(tmp_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(200, 10, text="Hello FoliaExchange test.",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    path = tmp_path / "sample.pdf"
    pdf.output(str(path))
    return str(path)


@pytest.fixture
def two_column_pdf(tmp_path):
    """A page with a left and right text column to test reading order."""
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)  # Letter
    page.insert_text((60, 100), "LEFT-1 first line of the left column.", fontsize=11)
    page.insert_text((60, 130), "LEFT-2 second line of the left column.", fontsize=11)
    page.insert_text((340, 100), "RIGHT-1 first line of the right column.", fontsize=11)
    page.insert_text((340, 130), "RIGHT-2 second line of the right column.", fontsize=11)
    path = tmp_path / "two_column.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture
def paper_pdf(tmp_path):
    """A paper-like page: full-width title + abstract above two columns."""
    import fitz
    doc = fitz.open()
    p = doc.new_page(width=612, height=792)
    p.insert_text((150, 70), "A Study of Document Conversion", fontsize=20)
    p.insert_text((72, 110), "Abstract spans the full width as one block.", fontsize=11)
    p.insert_text((60, 160), "Left A1 starts the left column here.", fontsize=11)
    p.insert_text((60, 185), "Left A2 continues the left column.", fontsize=11)
    p.insert_text((330, 160), "Right B1 starts the right column.", fontsize=11)
    p.insert_text((330, 185), "Right B2 continues the right column.", fontsize=11)
    path = tmp_path / "paper.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture
def table_pdf(tmp_path):
    """A page with a ruled 3x2 grid so find_tables() detects a real table."""
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=612, height=792)
    cells = [["Name", "Value"], ["Alpha", "1"], ["Beta", "2"]]
    x0, y0, cw, ch = 100, 200, 150, 30
    for r in range(3):
        for c in range(2):
            rx0, ry0 = x0 + c * cw, y0 + r * ch
            page.draw_rect(fitz.Rect(rx0, ry0, rx0 + cw, ry0 + ch),
                           color=(0, 0, 0), width=1)
            page.insert_text((rx0 + 5, ry0 + 20), cells[r][c], fontsize=11)
    path = tmp_path / "table.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture
def multi_page_pdf(tmp_path):
    """A two-page PDF to verify content flows without injected page breaks."""
    import fitz
    doc = fitz.open()
    for i in range(2):
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 100), f"Content on page {i + 1}.", fontsize=12)
    path = tmp_path / "multi_page.pdf"
    doc.save(str(path))
    doc.close()
    return str(path)


@pytest.fixture
def sample_docx(tmp_path):
    doc = Document()
    doc.add_heading("FoliaExchange Test", 0)
    doc.add_paragraph("Hello, this is a test document.")
    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return str(path)
