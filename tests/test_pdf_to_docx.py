import pytest
from pathlib import Path
from docx import Document
from converters.pdf_to_docx import convert


def _text(docx_path):
    return "\n".join(p.text for p in Document(docx_path).paragraphs)


def test_creates_docx_file(sample_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(sample_pdf, out)
    assert Path(out).exists()
    assert Path(out).stat().st_size > 1000


def test_extracts_text(sample_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(sample_pdf, out)
    assert "Hello FoliaExchange" in _text(out)


def test_two_column_reading_order(two_column_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(two_column_pdf, out)
    text = _text(out)
    # Left column must be read fully before the right column.
    assert text.index("LEFT-1") < text.index("LEFT-2") < text.index("RIGHT-1")


def test_paper_layout_order(paper_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(paper_pdf, out)
    t = _text(out)
    # Full-width title + abstract come before the columns; left column is read
    # fully before the right column.
    order = [t.index(k) for k in
             ("A Study", "Abstract", "Left A1", "Left A2", "Right B1", "Right B2")]
    assert order == sorted(order)


def test_reconstructs_real_tables(table_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(table_pdf, out)
    doc = Document(out)
    assert len(doc.tables) >= 1
    cells = [c.text.strip() for row in doc.tables[0].rows for c in row.cells]
    for expected in ("Name", "Value", "Alpha", "1", "Beta", "2"):
        assert expected in cells
    # Table text must NOT be duplicated as flowing paragraphs.
    assert "Alpha" not in _text(out)


def test_no_blank_pages_between_pages(multi_page_pdf, tmp_path):
    out = str(tmp_path / "output.docx")
    convert(multi_page_pdf, out)
    # Content flows: no per-page section/page breaks are inserted.
    assert Document(out).element.xpath(".//w:br[@w:type='page']") == []


def test_raises_on_missing_file(tmp_path):
    with pytest.raises(Exception):
        convert(str(tmp_path / "ghost.pdf"), str(tmp_path / "out.docx"))
